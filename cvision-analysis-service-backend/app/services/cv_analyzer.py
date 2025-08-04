# app/services/cv_analyzer.py
import os
import re
import json
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from loguru import logger
from sqlalchemy.orm import Session
from datetime import datetime

# Document processing
import PyPDF2
from docx import Document
import pdfplumber

# NLP and analysis
import spacy
import nltk
from textstat import flesch_reading_ease, flesch_kincaid_grade
from fuzzywuzzy import fuzz, process
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from app.models import CVFileModel, CVAnalysisResultModel, KeywordMatchModel, JobProfileModel
from app.core.config import settings
from app.core.constants import CVSections, MatchTypes, COMMON_SKILLS, EDUCATION_KEYWORDS, EXPERIENCE_KEYWORDS

class CVAnalyzer:
    """Main CV Analysis Service"""
    
    def __init__(self):
        self.nlp = None
        self.tfidf_vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
        self._load_nlp_model()
        self._download_nltk_data()
    
    def _load_nlp_model(self):
        """Load spaCy NLP model"""
        try:
            self.nlp = spacy.load(settings.SPACY_MODEL)
            logger.info(f"Loaded spaCy model: {settings.SPACY_MODEL}")
        except OSError:
            logger.warning(f"spaCy model {settings.SPACY_MODEL} not found. Install with: python -m spacy download {settings.SPACY_MODEL}")
            # Fallback to basic English model
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except OSError:
                logger.error("No spaCy model available. Please install en_core_web_sm")
                self.nlp = None
    
    def _download_nltk_data(self):
        """Download required NLTK data"""
        try:
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            nltk.download('averaged_perceptron_tagger', quiet=True)
        except Exception as e:
            logger.warning(f"Failed to download NLTK data: {e}")

    async def analyze_cv(self, cv_file: CVFileModel, db: Session) -> bool:
        """Main CV analysis method"""
        try:
            logger.info(f"Starting analysis for CV: {cv_file.FileName} (ID: {cv_file.Id})")
            
            # Update status to Processing
            cv_file.AnalysisStatus = "Processing"
            db.commit()
            
            # Extract text from CV
            extracted_text = self._extract_text_from_file(cv_file.FilePath, cv_file.FileType)
            if not extracted_text:
                raise ValueError("Failed to extract text from CV")
            
            # Store parsed text
            cv_file.ParsedText = extracted_text
            
            # Perform analysis
            analysis_result = self._analyze_cv_content(extracted_text, cv_file.Id)
            
            # Save analysis results to database
            self._save_analysis_results(cv_file, analysis_result, db)
            
            # Update CV status to Completed
            cv_file.AnalysisStatus = "Completed"
            cv_file.UpdatedAt = datetime.utcnow()
            db.commit()
            
            logger.info(f"Successfully completed analysis for CV: {cv_file.FileName}")
            return True
            
        except Exception as e:
            logger.error(f"Error analyzing CV {cv_file.FileName}: {e}")
            cv_file.AnalysisStatus = "Failed"
            cv_file.UpdatedAt = datetime.utcnow()
            db.commit()
            return False

    def _extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extract text from PDF or DOCX files"""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Remove dot if present and normalize to lowercase
            clean_file_type = file_type.lower().lstrip('.')
            
            if clean_file_type == "pdf":
                return self._extract_from_pdf(file_path)
            elif clean_file_type in ["docx", "doc"]:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber (more accurate than PyPDF2)"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # Fallback to PyPDF2 if pdfplumber fails
            if not text.strip():
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                        
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            
        return text.strip()

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return "\n".join(text)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""

    def _analyze_cv_content(self, text: str, cv_file_id: str) -> Dict[str, Any]:
        """Perform comprehensive CV analysis"""
        try:
            # Clean and preprocess text
            cleaned_text = self._clean_text(text)
            
            # Extract CV sections
            sections = self._extract_cv_sections(cleaned_text)
            
            # Analyze different aspects
            skills_analysis = self._analyze_skills(cleaned_text, sections)
            experience_analysis = self._analyze_experience(cleaned_text, sections)
            education_analysis = self._analyze_education(cleaned_text, sections)
            format_analysis = self._analyze_format(text, sections)
            
            # Calculate overall score
            overall_score = self._calculate_overall_score(
                skills_analysis, experience_analysis, education_analysis, format_analysis
            )
            
            # Find missing sections
            missing_sections = self._find_missing_sections(sections)
            
            # Find format issues
            format_issues = format_analysis.get('issues', [])
            
            return {
                'score': overall_score,
                'skills_analysis': skills_analysis,
                'experience_analysis': experience_analysis,
                'education_analysis': education_analysis,
                'format_analysis': format_analysis,
                'missing_sections': missing_sections,
                'format_issues': format_issues,
                'sections': sections
            }
            
        except Exception as e:
            logger.error(f"Error in CV content analysis: {e}")
            return {
                'score': 0,
                'missing_sections': ['analysis_error'],
                'format_issues': ['Failed to analyze CV content'],
                'sections': {}
            }

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep necessary punctuation
        text = re.sub(r'[^\w\s\-.,@()]+', '', text)
        return text.strip()

    def _extract_cv_sections(self, text: str) -> Dict[str, str]:
        """Extract different sections from CV"""
        sections = {}
        
        # Common section headers and their variations
        section_patterns = {
            CVSections.PERSONAL_INFO: [
                r'personal\s+information', r'contact\s+information', r'personal\s+details'
            ],
            CVSections.SUMMARY: [
                r'summary', r'profile', r'objective', r'about\s+me', r'career\s+objective'
            ],
            CVSections.EXPERIENCE: [
                r'experience', r'work\s+experience', r'professional\s+experience', 
                r'employment\s+history', r'career\s+history'
            ],
            CVSections.EDUCATION: [
                r'education', r'educational\s+background', r'academic\s+background',
                r'qualifications', r'academic\s+qualifications'
            ],
            CVSections.SKILLS: [
                r'skills', r'technical\s+skills', r'core\s+competencies', 
                r'key\s+skills', r'abilities'
            ],
            CVSections.CERTIFICATIONS: [
                r'certifications', r'certificates', r'training', r'courses'
            ],
            CVSections.PROJECTS: [
                r'projects', r'key\s+projects', r'notable\s+projects'
            ],
            CVSections.LANGUAGES: [
                r'languages', r'language\s+skills'
            ]
        }
        
        text_lower = text.lower()
        lines = text.split('\n')
        
        for section_name, patterns in section_patterns.items():
            for pattern in patterns:
                matches = list(re.finditer(pattern, text_lower))
                if matches:
                    # Find the section content
                    start_pos = matches[0].end()
                    
                    # Find next section or end of text
                    next_section_pos = len(text)
                    for other_section, other_patterns in section_patterns.items():
                        if other_section != section_name:
                            for other_pattern in other_patterns:
                                other_matches = list(re.finditer(other_pattern, text_lower[start_pos:]))
                                if other_matches:
                                    candidate_pos = start_pos + other_matches[0].start()
                                    if candidate_pos < next_section_pos:
                                        next_section_pos = candidate_pos
                    
                    section_content = text[start_pos:next_section_pos].strip()
                    sections[section_name] = section_content
                    break
        
        return sections

    def _analyze_skills(self, text: str, sections: Dict[str, str]) -> Dict[str, Any]:
        """Analyze skills mentioned in CV"""
        skills_text = sections.get(CVSections.SKILLS, text)
        found_skills = []
        skill_matches = []
        
        text_lower = text.lower()
        
        # Enhanced skill detection with exact and fuzzy matching
        for skill in COMMON_SKILLS:
            skill_lower = skill.lower()
            
            # Direct exact match
            if skill_lower in text_lower:
                found_skills.append(skill)
                skill_matches.append({
                    'keyword': skill,
                    'confidence': 1.0,
                    'match_type': MatchTypes.EXACT
                })
                continue
            
            # Handle special cases for compound skills
            if '.' in skill_lower:
                # For skills like ".net core", check variations
                variations = [
                    skill_lower,
                    skill_lower.replace('.', ''),  # "net core" 
                    skill_lower.replace(' ', ''),  # ".netcore"
                    skill_lower.replace('.', '').replace(' ', '')  # "netcore"
                ]
                for variation in variations:
                    if variation in text_lower:
                        found_skills.append(skill)
                        skill_matches.append({
                            'keyword': skill,
                            'confidence': 0.9,
                            'match_type': MatchTypes.FUZZY
                        })
                        break
        
        # Enhanced pattern matching for specific technologies
        enhanced_patterns = [
            # Programming Languages
            r'\b(python|java|javascript|typescript|c#|c\+\+|php|ruby|go|rust|swift|kotlin|c|scala|dart)\b',
            # .NET Technologies  
            r'\b(\.?net\s*core?|asp\.?net|entity\s*framework|blazor|mvc|webapi|signalr)\b',
            # Frontend Frameworks
            r'\b(react|angular|vue\.?js|next\.?js|nuxt\.?js|svelte|ember)\b',
            # Backend Frameworks
            r'\b(django|fastapi|flask|spring|laravel|rails|express\.?js)\b',
            # Databases
            r'\b(postgresql|mysql|mongodb|redis|elasticsearch|mssql|sql\s*server|sqlite|oracle)\b',
            # DevOps & Tools
            r'\b(docker|kubernetes|jenkins|git|github|gitlab|ci/cd|nginx|apache|rabbitmq)\b',
            # Cloud Platforms
            r'\b(aws|azure|gcp|google\s*cloud|amazon\s*web\s*services|heroku|vercel)\b',
            # Styling & Design
            r'\b(tailwind|bootstrap|sass|scss|less|css3|html5)\b'
        ]
        
        for pattern in enhanced_patterns:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                skill = match.group().strip()
                # Normalize the skill name
                skill = skill.replace('.', '').replace('js', ' JS').strip()
                if skill not in [s['keyword'] for s in skill_matches]:
                    skill_matches.append({
                        'keyword': skill,
                        'confidence': 1.0,
                        'match_type': MatchTypes.EXACT
                    })
        
        # Calculate enhanced skills score
        unique_skills = list(set(found_skills + [s['keyword'] for s in skill_matches]))
        skills_score = min(len(unique_skills) * 3, 100)  # More realistic scoring
        
        return {
            'found_skills': unique_skills,
            'skill_matches': skill_matches,
            'skills_count': len(unique_skills),
            'skills_score': skills_score
        }

    def _analyze_experience(self, text: str, sections: Dict[str, str]) -> Dict[str, Any]:
        """Analyze work experience"""
        experience_text = sections.get(CVSections.EXPERIENCE, text)
        
        # Enhanced patterns for extracting years of experience
        year_patterns = [
            r'(\d+)\s*\+?\s*years?\s+(?:of\s+)?experience',
            r'experience\s+(?:of\s+)?(\d+)\s*\+?\s*years?',
            r'(\d+)\s*\+?\s*years?\s+(?:in|with)',
            r'(\d+)\s*\+?\s*year\s+(?:of\s+)?(?:experience|work)',
        ]
        
        years_found = []
        for pattern in year_patterns:
            matches = re.finditer(pattern, text.lower())
            for match in matches:
                years_found.append(int(match.group(1)))
        
        # Enhanced job position detection
        job_titles = [
            r'(?:full-stack|front-end|back-end|senior|junior|lead|principal)?\s*(?:software\s+)?(?:engineer|developer|programmer)',
            r'(?:software\s+)?(?:architect|analyst|consultant|specialist)',
            r'(?:project\s+)?(?:manager|director|coordinator|lead)',
            r'(?:data\s+)?(?:scientist|analyst|engineer)',
            r'(?:devops|system\s+administrator|it\s+specialist)',
            r'intern(?:ship)?'
        ]
        
        job_count = 0
        for pattern in job_titles:
            matches = re.finditer(pattern, text.lower())
            job_count += len(list(matches))
        
        # Enhanced date range detection for calculating experience duration
        date_patterns = [
            r'(\w+\s+\d{4})\s*[–\-−—]\s*(\w+\s+\d{4}|present|current)',
            r'(\d{4})\s*[–\-−—]\s*(\d{4}|present|current)',
            r'(\w+\s+\d{4})\s*[–\-−—]\s*(present|current)',
            r'(august\s+2024)\s*[–\-−—]\s*(january\s+2025)',
            r'(january\s+2024)\s*[–\-−—]\s*(february\s+2024)',
            r'(january\s+2022)\s*[–\-−—]\s*(present|current)',
        ]
        
        work_periods = []
        for pattern in date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                start_date = match.group(1)
                end_date = match.group(2)
                work_periods.append((start_date, end_date))
        
        # Calculate total experience from date ranges
        calculated_years = 0
        current_year = 2025  # Current year
        current_month = 8   # Current month (August)
        
        # Month name to number mapping
        months = {
            'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
            'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
        }
        
        for start_date, end_date in work_periods:
            try:
                # Extract start year and month
                start_year_match = re.search(r'\d{4}', start_date)
                start_month_match = re.search(r'(january|february|march|april|may|june|july|august|september|october|november|december)', start_date.lower())
                
                if start_year_match:
                    start_year = int(start_year_match.group())
                    start_month = months.get(start_month_match.group(), 1) if start_month_match else 1
                    
                    # Extract end year and month
                    if end_date.lower() in ['present', 'current']:
                        end_year = current_year
                        end_month = current_month
                    else:
                        end_year_match = re.search(r'\d{4}', end_date)
                        end_month_match = re.search(r'(january|february|march|april|may|june|july|august|september|october|november|december)', end_date.lower())
                        
                        if end_year_match:
                            end_year = int(end_year_match.group())
                            end_month = months.get(end_month_match.group(), 12) if end_month_match else 12
                        else:
                            end_year = current_year
                            end_month = current_month
                    
                    # Calculate total months and convert to years
                    total_months = (end_year - start_year) * 12 + (end_month - start_month)
                    period_years = max(0, total_months / 12)
                    calculated_years += period_years
                    
            except (ValueError, AttributeError):
                continue
        
        # Use the maximum of explicitly mentioned years or calculated years
        total_years = max(max(years_found) if years_found else 0, calculated_years)
        
        # Enhanced scoring
        experience_score = 0
        if total_years > 0:
            experience_score += min(total_years * 15, 70)  # Up to 70 points for years
        if job_count > 0:
            experience_score += min(job_count * 10, 30)   # Up to 30 points for positions
        
        experience_score = min(experience_score, 100)
        
        return {
            'years_of_experience': total_years,
            'job_positions_count': job_count,
            'work_periods': work_periods,
            'experience_score': experience_score
        }

    def _analyze_education(self, text: str, sections: Dict[str, str]) -> Dict[str, Any]:
        """Analyze educational background"""
        education_text = sections.get(CVSections.EDUCATION, text)
        
        degrees_found = []
        institutions_found = []
        
        # Enhanced degree patterns
        degree_patterns = [
            # Bachelor's degrees
            r'\b(?:bachelor[\'s]*|bs|ba|bsc|be|btech|b\.?\s*tech|b\.?\s*sc|b\.?\s*a)\s*(?:degree|of)?\s*(?:in|of)?\s*\w*',
            # Master's degrees  
            r'\b(?:master[\'s]*|ms|ma|msc|me|mtech|mba|m\.?\s*tech|m\.?\s*sc|m\.?\s*a)\s*(?:degree|of)?\s*(?:in|of)?\s*\w*',
            # PhD and Doctorate
            r'\b(?:phd|ph\.?\s*d\.?|doctorate|doctoral)\s*(?:degree|of)?\s*(?:in|of)?\s*\w*',
            # Associate and other degrees
            r'\b(?:associate|diploma|certificate)\s*(?:degree|of)?\s*(?:in|of)?\s*\w*',
            # Specific degree mentions
            r'\b(?:economics|computer\s*science|information\s*management|engineering|business)\s*(?:bachelor|master|degree)',
        ]
        
        for pattern in degree_patterns:
            matches = re.finditer(pattern, text.lower(), re.IGNORECASE)
            for match in matches:
                degree = match.group().strip()
                if degree and degree not in degrees_found:
                    degrees_found.append(degree)
        
        # Enhanced institution patterns
        institution_patterns = [
            # University patterns with specific names
            r'\b(?:anadolu|dumlup[iı]nar|\w+)\s+university\b',
            r'\b42\s+kocaeli\b',
            r'\b(?:école|ecole)\s+42\b',
            r'\buniversity\s+of\s+\w+\b',
            # College patterns
            r'\b\w+\s+(?:college|institut[eo]?|school)\b',
            # Technical schools
            r'\b(?:technical\s+)?(?:institute|academy|school)\s+of\s+\w+\b',
        ]
        
        for pattern in institution_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                institution = match.group().strip()
                if institution and institution not in institutions_found:
                    institutions_found.append(institution)
        
        # Extract graduation years and calculate recency bonus
        graduation_years = []
        year_patterns = [
            r'(?:graduated|graduation|completed).*?(\d{4})',
            r'(\d{4})\s*[-–]\s*(?:\d{4}|present|current)',
            r'(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+(\d{4})'
        ]
        
        for pattern in year_patterns:
            matches = re.finditer(pattern, text.lower())
            for match in matches:
                year = int(match.group(1))
                if 1980 <= year <= 2025:  # Reasonable year range
                    graduation_years.append(year)
        
        # Calculate enhanced education score
        education_score = 0
        
        # Base score for degrees
        degree_scores = {
            'phd': 40, 'doctorate': 40, 'doctoral': 40,
            'master': 30, 'mba': 35, 'ms': 30, 'ma': 30, 'msc': 30,
            'bachelor': 25, 'bs': 25, 'ba': 25, 'bsc': 25, 'btech': 25,
            'associate': 15, 'diploma': 15, 'certificate': 10
        }
        
        max_degree_score = 0
        for degree in degrees_found:
            for degree_type, score in degree_scores.items():
                if degree_type in degree.lower():
                    max_degree_score = max(max_degree_score, score)
                    break
        
        education_score += max_degree_score
        
        # Bonus for multiple degrees
        if len(degrees_found) > 1:
            education_score += min(len(degrees_found) * 5, 20)
        
        # Bonus for recognized institutions
        if institutions_found:
            education_score += min(len(institutions_found) * 10, 30)
        
        # Recency bonus (more recent education gets bonus)
        if graduation_years:
            recent_year = max(graduation_years)
            if recent_year >= 2020:
                education_score += 10
            elif recent_year >= 2015:
                education_score += 5
        
        education_score = min(education_score, 100)
        
        return {
            'degrees_found': degrees_found,
            'institutions_found': institutions_found,
            'graduation_years': graduation_years,
            'education_score': education_score
        }

    def _analyze_format(self, text: str, sections: Dict[str, str]) -> Dict[str, Any]:
        """Analyze CV format and structure"""
        issues = []
        
        # Enhanced readability check
        readability_score = flesch_reading_ease(text)
        if readability_score < 20:
            issues.append("CV text is very difficult to read")
        elif readability_score < 40:
            issues.append("CV text could be more readable")
        
        # Check length
        word_count = len(text.split())
        if word_count < 150:
            issues.append("CV is too short (less than 150 words)")
        elif word_count > 3000:
            issues.append("CV is too long (more than 3000 words)")
        
        # Enhanced contact information patterns
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_patterns = [
            r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+?\d{2,3}\s?\d{3}\s?\d{3}\s?\d{2}\s?\d{2}',  # International format
            r'\(\+\d{2}\)\s?\d{3}\s?\d{3}\s?\d{2}\s?\d{2}'   # Turkish format
        ]
        
        has_email = bool(re.search(email_pattern, text))
        has_phone = any(re.search(pattern, text) for pattern in phone_patterns)
        
        if not has_email:
            issues.append("No email address found")
        if not has_phone:
            issues.append("No phone number found")
        
        # Check for professional links
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        github_pattern = r'github\.com/[\w-]+'
        
        has_linkedin = bool(re.search(linkedin_pattern, text, re.IGNORECASE))
        has_github = bool(re.search(github_pattern, text, re.IGNORECASE))
        
        # Check for section headers (structure quality)
        expected_sections = ['experience', 'education', 'skills', 'summary', 'projects']
        found_section_headers = 0
        
        for section in expected_sections:
            if re.search(rf'\b{section}\b', text, re.IGNORECASE):
                found_section_headers += 1
        
        if found_section_headers < 3:
            issues.append("Missing important CV sections")
        
        # Check for dates (shows experience timeline)
        date_patterns = [
            r'\d{4}\s*[-–]\s*(?:\d{4}|present|current)',
            r'(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{4}'
        ]
        
        has_dates = any(re.search(pattern, text, re.IGNORECASE) for pattern in date_patterns)
        if not has_dates:
            issues.append("No date information found in work experience")
        
        # Check for bullet points or structure indicators
        structure_indicators = ['-', '•', '◦', '▪', '▫', '○', '●']
        has_structure = any(indicator in text for indicator in structure_indicators)
        
        # Calculate enhanced format score
        format_score = 100
        
        # Penalties for issues
        format_score -= len(issues) * 12
        
        # Bonuses for good formatting
        if has_linkedin:
            format_score += 5
        if has_github:
            format_score += 5
        if found_section_headers >= 4:
            format_score += 10
        if has_structure:
            format_score += 5
        if 200 <= word_count <= 1500:  # Optimal length
            format_score += 10
        if readability_score >= 60:  # Good readability
            format_score += 10
        
        format_score = max(min(format_score, 100), 0)
        
        return {
            'readability_score': readability_score,
            'word_count': word_count,
            'has_email': has_email,
            'has_phone': has_phone,
            'has_linkedin': has_linkedin,
            'has_github': has_github,
            'section_headers_found': found_section_headers,
            'has_dates': has_dates,
            'has_structure': has_structure,
            'issues': issues,
            'format_score': format_score
        }

    def _find_missing_sections(self, sections: Dict[str, str]) -> List[str]:
        """Find missing important sections"""
        required_sections = [
            CVSections.EXPERIENCE,
            CVSections.EDUCATION,
            CVSections.SKILLS
        ]
        
        missing = []
        for section in required_sections:
            if section not in sections or not sections[section].strip():
                missing.append(section)
        
        return missing

    def _calculate_overall_score(self, skills_analysis: Dict, experience_analysis: Dict, 
                               education_analysis: Dict, format_analysis: Dict) -> int:
        """Calculate weighted overall score"""
        skills_score = skills_analysis.get('skills_score', 0)
        experience_score = experience_analysis.get('experience_score', 0)
        education_score = education_analysis.get('education_score', 0)
        format_score = format_analysis.get('format_score', 0)
        
        overall_score = (
            skills_score * settings.SKILLS_WEIGHT +
            experience_score * settings.EXPERIENCE_WEIGHT +
            education_score * settings.EDUCATION_WEIGHT +
            format_score * settings.FORMAT_WEIGHT
        )
        
        return int(min(overall_score, 100))

    def _save_analysis_results(self, cv_file: CVFileModel, analysis_result: Dict[str, Any], db: Session):
        """Save analysis results to database"""
        try:
            # Create or update analysis result
            existing_result = db.query(CVAnalysisResultModel).filter(
                CVAnalysisResultModel.CVFileId == cv_file.Id
            ).first()
            
            if existing_result:
                # Update existing result
                result = existing_result
                result.UpdatedAt = datetime.utcnow()
            else:
                # Create new result
                result = CVAnalysisResultModel(CVFileId=cv_file.Id)
            
            # Set analysis data
            result.Score = analysis_result['score']
            result.MissingSections = analysis_result['missing_sections']
            result.FormatIssues = analysis_result['format_issues']
            
            if not existing_result:
                db.add(result)
            db.flush()  # Get the ID
            
            # Save keyword matches
            self._save_keyword_matches(result.Id, analysis_result.get('skills_analysis', {}), db)
            
            db.commit()
            logger.info(f"Saved analysis results for CV file: {cv_file.Id}")
            
        except Exception as e:
            logger.error(f"Error saving analysis results: {e}")
            db.rollback()
            raise

    def _save_keyword_matches(self, analysis_result_id: str, skills_analysis: Dict[str, Any], db: Session):
        """Save keyword matches to database"""
        try:
            # Delete existing keyword matches
            db.query(KeywordMatchModel).filter(
                KeywordMatchModel.CVAnalysisResultId == analysis_result_id
            ).delete()
            
            # Add new keyword matches
            skill_matches = skills_analysis.get('skill_matches', [])
            for match in skill_matches:
                # Calculate relevance score based on confidence and match type
                relevance_score = int(match.get('confidence', 1.0) * 100)
                
                keyword_match = KeywordMatchModel(
                    CVAnalysisResultId=analysis_result_id,
                    Keyword=match['keyword'],
                    IsMatched=True,
                    Count=1,  # Count of occurrences
                    MatchCount=1,  # Keep for backward compatibility
                    Relevance=relevance_score  # Relevance score (0-100)
                )
                db.add(keyword_match)
            
            db.flush()
            
        except Exception as e:
            logger.error(f"Error saving keyword matches: {e}")
            raise